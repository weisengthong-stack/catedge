#!/usr/bin/env python3
"""
KYC Document Processor - Streamlit UI
A beautiful interface for processing KYC documents with complete Word generation
"""

import streamlit as st
import os
import json
import tempfile
from datetime import datetime
from pathlib import Path
from typing import List, Optional
import io

# Docling imports
from docling.document_converter import DocumentConverter

# LlamaExtract imports
from pydantic import BaseModel, Field
from llama_cloud_services import LlamaExtract

# Word document generation imports
import pandas as pd
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Set API key

NBSP = "\u00A0"

# ========== DATA MODELS ==========
class CustomerProfile(BaseModel):
    entityType: str = Field(description="Type of entity, e.g., ASSOCIATION.")
    corporateWebsite: Optional[str] = Field(description="URL of the corporate website, if available.")
    fatfjurisdiction: Optional[str] = Field(description="FATF jurisdiction associated with the customer.")
    industry: str = Field(description="Industry sector of the customer's business, e.g., AGRICULTURE.")
    onBoardingMode: Optional[str] = Field(description="Mode of onboarding the customer, e.g., FACE-TO-FACE.")
    ownershipStructureLayer: Optional[str] = Field(description="Layer of the ownership structure.")
    paymentMode: Optional[List[str]] = Field(description="Preferred payment modes of the customer.")
    productServiceComplexity: str = Field(description='Complexity level of the products or services offered. Filled in by either "SIMPLE" or "COMPLEX"')
    sourceOfFunds: Optional[str] = Field(description="Source of funds for the customer's transactions.")
    natureOfBusinessRelationship: Optional[str] = Field(description="Nature of the business relationship with the customer.")
    bankAccount: Optional[List[str]] = Field(description="List of bank accounts associated with the customer.")
    additionalInformation: Optional[str] = Field(description="Any additional information about the customer.")
    incorporated: bool = Field(description="Indicates whether the customer is incorporated.")
    name: str = Field(description="Official name of the customer.")
    alias: Optional[str] = Field(description="List of aliases or alternative names used by the customer.")
    formerName: Optional[str] = Field(description="List of former names of the customer.")
    countryOfIncorporation: str = Field(description="Country where the customer was incorporated.")
    countryOfOperation: str = Field(description="List of countries where the customer operates.")
    address: str = Field(description='List of addresses associated with the customer. if exists, take the "Address of principal place of business"')
    incorporateNumber: str = Field(description="Incorporation number of the customer.")
    phone: Optional[str] = Field(description="List of phone numbers associated with the customer.")
    email: Optional[str] = Field(description="List of email addresses associated with the customer.")
    dateOfIncorporation: Optional[str] = Field(description="Date when the customer was incorporated (format: YYYY-MM-DD).")
    dateOfIncorporationExpiry: Optional[str] = Field(description="Expiry date of incorporation (format: YYYY-MM-DD).")
    imonumber: Optional[str] = Field(description="IMO number of the customer, if applicable.")
    active: bool = Field(description="Indicates whether the customer profile is currently active.")
    profileReferenceId: str = Field(description="Unique identifier for the customer profile.")
    other: Optional[str] = Field(description="Contains additional information about the customer, such as entity type, industry, and payment preferences.")
    particular: Optional[str] = Field(description="Specific details about the customer, such as incorporation status, name, and contact information.")
    type: Optional[str] = Field(description="Type of customer, e.g., CORPORATE.")
    domainId: Optional[List[str]] = Field(description="List of domain IDs associated with the customer.")
    assigneeId: Optional[str] = Field(description="ID of the assignee responsible for the customer.")

class ArtemisProfile(BaseModel):
    company_name: str = Field(description="Name of the company. Extracted from the 'CORPORATE INFORMATION' section.")
    business_reg_no: str = Field(description="Business registration number from 'Incorporation Number' field in 'CORPORATE INFORMATION' section.")
    pep_status: bool = Field(description="The latest PEP status from the 'SCREENING & SEARCH CONCLUSION' section for the latest date.")
    client_alert_list: bool = Field(description="Indicates 'Yes' if there is any World Check matches where 'MATCH TYPE' is 'TRUE HIT' or 'MATCH STRENGTH' is 'EXACT'")
    client_investigation: bool = Field(description="Indicates 'Yes' if there is any World Check matches where 'MATCH TYPE' is 'TRUE HIT' or 'MATCH STRENGTH' is 'EXACT'")
    internet_search_alerts: bool = Field(description="Indicates 'Yes' if there is any of Internet Search matches where 'MATCH TYPE' is 'TRUE HIT' or 'MATCH STRENGTH' is 'EXACT'")
    any_investigation: bool = Field(description="Indicates 'Yes' if there is any 'YES' in the 'SCREENING & SEARCH CONCLUSION' part for 'SANCTIONS' and 'ADVERSE NEWS', otherwise 'No'.")
    any_adverse_news: bool = Field(description="Indicates 'Yes' if there is any 'YES' in the 'SCREENING & SEARCH CONCLUSION' part for 'ADVERSE NEWS', otherwise 'No'.")

# ========== HELPER FUNCTIONS ==========
def safe_get(d, path, default="-"):
    cur = d
    for p in path:
        if isinstance(cur, dict):
            cur = cur.get(p)
        else:
            return default
    return cur if cur not in (None, "", [], {}) else default

def load_excel_kv(excel_path):
    kv = {}
    try:
        xls = pd.read_excel(excel_path, sheet_name=None, header=None, dtype=str, engine="openpyxl")
    except Exception as e:
        st.warning(f"Error reading Excel: {e}")
        return kv
    
    for _, df in xls.items():
        df = df.fillna("")
        rows, cols = df.shape
        for r in range(rows):
            for c in range(cols - 1):
                key = str(df.iat[r, c]).strip()
                val = str(df.iat[r, c+1]).strip()
                if key and key.lower() != "nan" and val and val.lower() != "nan":
                    if " " in key or ":" in key:
                        kv[key.lower()] = val
    return kv

def find_excel_value(kv, *labels, default="-"):
    for lab in labels:
        L = lab.lower()
        if L in kv and kv[L]:
            return kv[L]
        for k, v in kv.items():
            if L in k and v:
                return v
    return default

def yn(val, dash="-"):
    if val in (True, "True", "true", "YES", "Yes", "yes", "Y", "y", 1):
        return "Yes"
    if val in (False, "False", "false", "NO", "No", "no", "N", "n", 0):
        return "No"
    return dash

# Word formatting helpers
def set_margins(section, top=1.8, bottom=1.5, left=1.8, right=1.8):
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)

def add_header_footer(document):
    hdr = document.sections[0].header
    p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    run = p.add_run("CONFIDENTIAL")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run.font.size = Pt(8)

    ftr = document.sections[0].footer
    p2 = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run("KYC Form Pg ")
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    r = OxmlElement('w:r'); t = OxmlElement('w:t'); t.text = ''
    r.append(t); fld.append(r)
    p2._element.append(fld)
    p2.add_run("/Aug 2025")
    for r in p2.runs:
        r.font.size = Pt(8)

def add_para(document, text, size=9.5, bold=False, italic=False, space_after=0.8):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after*12/0.846)
    return p

def make_table(document, rows, cols, col_widths_cm=None, align=WD_TABLE_ALIGNMENT.LEFT):
    tbl = document.add_table(rows=rows, cols=cols)
    tbl.alignment = align
    tbl.autofit = False
    if col_widths_cm:
        for i, w in enumerate(col_widths_cm):
            if i < len(tbl.columns):
                for cell in tbl.columns[i].cells:
                    cell.width = Cm(w)
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            vAlign = OxmlElement('w:vAlign')
            vAlign.set(qn('w:val'), "top")
            tcPr.append(vAlign)
            set_cell_borders(cell)
    return tbl

def set_cell_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '6')
        elem.set(qn('w:space'), '0')
        elem.set(qn('w:color'), '000000')
        tcBorders.append(elem)
    tcPr.append(tcBorders)

def shade_row(row):
    for cell in row.cells:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        cell._tc.get_or_add_tcPr().append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

def set_cell_text(cell, text, size=9.5, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text if (text not in (None, "")) else "-")
    run.font.size = Pt(size)
    run.bold = bold

def style_title(document, text, size=16):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)

# ========== PROCESSING FUNCTIONS ==========
def convert_pdf_to_markdown(pdf_bytes, filename):
    """Convert PDF to markdown"""
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf', mode='wb') as tmp:
        tmp.write(pdf_bytes)
        tmp_path = tmp.name
    
    try:
        converter = DocumentConverter()
        result = converter.convert(tmp_path)
        markdown_output = result.document.export_to_markdown()
        doc_name = Path(filename).stem
        return markdown_output, doc_name
    except Exception as e:
        st.error(f"Error converting {filename}: {str(e)}")
        raise
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

def extract_customer_profile(markdown_text):
    """Extract customer profile"""
    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.md', encoding='utf-8') as tmp:
        tmp.write(markdown_text)
        tmp_path = tmp.name
    
    try:
        extractor = LlamaExtract()
        agent = extractor.get_agent(name="customer-profile-parser")
        customer_data = agent.extract(tmp_path)
        return customer_data.data
    except Exception as e:
        st.error(f"Error extracting customer profile: {str(e)}")
        raise
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

def extract_artemis_profile(markdown_text):
    """Extract Artemis profile"""
    with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.md', encoding='utf-8') as tmp:
        tmp.write(markdown_text)
        tmp_path = tmp.name
    
    try:
        extractor = LlamaExtract()
        agent = extractor.get_agent(name="artermis-extractor")
        artemis_data = agent.extract(tmp_path)
        return artemis_data.data
    except Exception as e:
        st.error(f"Error extracting Artemis profile: {str(e)}")
        raise
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

def generate_word_document(j1, j2, excel_bytes):
    """Generate complete Word document from data (all 7 pages)"""
    # Save Excel to temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx', mode='wb') as tmp:
        tmp.write(excel_bytes)
        excel_path = tmp.name
    
    try:
        kv = load_excel_kv(excel_path)
    except Exception as e:
        st.error(f"Error reading Excel file: {str(e)}")
        raise
    finally:
        if os.path.exists(excel_path):
            os.unlink(excel_path)
    
    # Extract key information
    client_name = safe_get(j1, ["particular", "name"], "-")
    year_ending = find_excel_value(kv, "Year Ending", "Year Ending (if applicable)", default="-")
    nature_of_service = find_excel_value(kv, "Nature of Service", "Nature of service / Purpose of transaction", default="-")
    kyc_risk_lookup = find_excel_value(kv, "KYC Risk Rating", default="")
    kyc_risk = kyc_risk_lookup.title() if kyc_risk_lookup else "-"
    
    biz_reg_no = safe_get(j2, ["business_reg_no"], "-") or safe_get(j1, ["particular", "incorporateNumber"], "-")
    address = safe_get(j1, ["particular", "address"], "-")
    senior_mgmt = find_excel_value(kv, "Senior Management", "Directors", default="-").replace(",", ", ")
    
    type_of_engagement = find_excel_value(kv, "Type of engagement", "Type of engagement (Recurring/Non-recurring)", default="-")
    proposed_services = find_excel_value(kv, "Proposed services", default="-")
    office = find_excel_value(kv, "Which office will accept/retain the client?", default="-")
    
    pep = yn(safe_get(j2, ["pep_status"], False))
    alert_list = yn(safe_get(j2, ["client_alert_list"], False))
    investigation = yn(safe_get(j2, ["client_investigation"], False) or safe_get(j2, ["any_investigation"], False))
    adverse = yn(safe_get(j2, ["internet_search_alerts"], False) or safe_get(j2, ["any_adverse_news"], False))
    
    # Create document
    document = Document()
    set_margins(document.sections[0], 1.8, 1.5, 1.8, 1.8)
    add_header_footer(document)
    
    # ========== PAGE 1 ==========
    style_title(document, "Know Your Client ('KYC') Assessment Form")
    add_para(document, "(For Non-Gazetted Activities)", size=9)
    
    tbl = make_table(document, 5, 2, [7.5, 10.0])
    shade_row(tbl.rows[0])
    set_cell_text(tbl.cell(0,0), "Details", bold=True)
    set_cell_text(tbl.cell(0,1), "")
    set_cell_text(tbl.cell(1,0), "Name of Client")
    set_cell_text(tbl.cell(1,1), client_name)
    set_cell_text(tbl.cell(2,0), "Year Ending (if applicable)")
    set_cell_text(tbl.cell(2,1), year_ending)
    set_cell_text(tbl.cell(3,0), "Nature of Service /\nPurpose of transaction")
    set_cell_text(tbl.cell(3,1), nature_of_service)
    set_cell_text(tbl.cell(4,0), "KYC Risk Rating")
    set_cell_text(tbl.cell(4,1), kyc_risk)
    
    document.add_paragraph()
    
    tbl = make_table(document, 6, 3, [9.5, 2.0, 6.0])
    shade_row(tbl.rows[0])
    set_cell_text(tbl.cell(0,0), "Summary of Relevant Information", bold=True)
    set_cell_text(tbl.cell(0,1), "Yes / No", bold=True)
    set_cell_text(tbl.cell(0,2), "Details", bold=True)
    
    rows_data = [
        ("Politically Exposed Person ('PEP')\n(only if the person is a PEP within the\nlast 12 months preceding the Artemis search dated [ ])", pep, "(State the name of the PEP)"),
        ("Higher Risk\nJurisdiction(s)/country(ies)", "No", "(State the name of the jurisdiction/county)"),
        ("Client/BO/Director listed in alert list\nissued by authorities", alert_list, "(State the Client/BO/Director's name and the alert list's name)"),
        ("Client/BO/Director under\ninvestigation orders issued by\nauthorities", investigation, "(State the Client/BO/Director's name and the authorities' name)"),
        ("Adverse information about the\nClient/BO/Director", adverse, "(State the Client/BO/Director's name and reference of the news in the KYC)"),
    ]
    for i, (a,b,c) in enumerate(rows_data, start=1):
        set_cell_text(tbl.cell(i,0), a)
        set_cell_text(tbl.cell(i,1), b)
        set_cell_text(tbl.cell(i,2), c)
    
    document.add_paragraph()
    add_para(document, "Objective", size=11, bold=True, space_after=0.2)
    add_para(document, "To document our KYC risk assessment and response during the client acceptance / retention stage", size=9.5)
    add_para(document, "Completion of KYC Assessment (Refer to the KYC Guidance Notes)", size=9)
    
    document.add_page_break()
    
    # ========== PAGE 2 ==========
    add_para(document, "KYC Assessment Form (continued)", size=14, bold=True)
    add_para(document, "Part A – Client Due Diligence", size=12, bold=True)
    
    tbl = make_table(document, 3, 5, [3.0, 4.5, 3.5, 2.5, 2.5])
    shade_row(tbl.rows[0])
    for c,h in enumerate(["Tasks", "Name", "Designation", "Signature", "Dates"]):
        set_cell_text(tbl.cell(0,c), h, bold=True)
    row_data = [
        ["Prepared by", find_excel_value(kv, "Prepared by", default="-"),
         find_excel_value(kv, "Prepared by Designation", default="-"),
         find_excel_value(kv, "Prepared by Signature", default="-"),
         find_excel_value(kv, "Prepared by Date", default="-")],
        ["Reviewed by (Engagement team)", find_excel_value(kv, "Reviewed by Name", default="-"),
         find_excel_value(kv, "Reviewed by Designation", default="-"),
         find_excel_value(kv, "Reviewed by Signature", default="-"),
         find_excel_value(kv, "Reviewed by Date", default="-")],
    ]
    for r, row in enumerate(row_data, start=1):
        for c, val in enumerate(row):
            set_cell_text(tbl.cell(r,c), val)
    
    document.add_paragraph()
    
    # Approved by
    tbl = make_table(document, 2, 3, [6.5, 4.5, 3.5])
    shade_row(tbl.rows[0])
    for c, h in enumerate(["Approved by","Designation","Date"]):
        set_cell_text(tbl.cell(0,c), h, bold=True)
    vals = [find_excel_value(kv, "Approved by Name", default="-"),
            find_excel_value(kv, "Approved by Designation", default="-"),
            find_excel_value(kv, "Approved by Date", default="-")]
    for c,v in enumerate(vals): set_cell_text(tbl.cell(1,c), v)
    
    document.add_paragraph()
    
    # Reviewed by (Deputy Compliance Officer)
    tbl = make_table(document, 2, 3, [6.5, 4.5, 3.5])
    shade_row(tbl.rows[0])
    for c, h in enumerate(["Reviewed by","Designation","Date"]):
        set_cell_text(tbl.cell(0,c), h, bold=True)
    vals = [find_excel_value(kv, "Reviewed2 Name", default="-"),
            find_excel_value(kv, "Reviewed2 Designation", default="-"),
            find_excel_value(kv, "Reviewed2 Date", default="-")]
    for c,v in enumerate(vals): set_cell_text(tbl.cell(1,c), v)
    
    document.add_paragraph()
    
    # Engagement Approved by
    tbl = make_table(document, 2, 3, [6.5, 4.5, 3.5])
    shade_row(tbl.rows[0])
    for c, h in enumerate(["Engagement Approved by","Designation","Date"]):
        set_cell_text(tbl.cell(0,c), h, bold=True)
    vals = [find_excel_value(kv, "Engagement Approved Name", default="-"),
            find_excel_value(kv, "Engagement Approved Designation", default="-"),
            find_excel_value(kv, "Engagement Approved Date", default="-")]
    for c,v in enumerate(vals): set_cell_text(tbl.cell(1,c), v)
    
    document.add_page_break()
    
    # ========== PAGE 3 ==========
    add_para(document, "KYC Assessment Form (continued)", size=14, bold=True)
    add_para(document, "(For Non-Gazetted Activities)", size=9)
    add_para(document, "Part A – Client Due Diligence", size=12, bold=True)
    add_para(document, "Section 1 – Engagement Information", size=11, bold=True)
    
    rows_sec1 = [
        ["1.", "Type of client", "Individual (Natural Person) ('IND')"],
        ["",   "",               "Legal person ('LP')"],
        ["",   "",               "Legal arrangement('LA')"],
        ["",   "",               "Club, Society and Charity('CSC')"],
        ["2.", "Type of engagement\n(Recurring/Non-recurring)", type_of_engagement],
        ["3.", "Nature of Service\n(Audit/Tax/Advisory/BSO)", nature_of_service],
        ["4.", "Proposed services", proposed_services],
        ["5.", "Which office will accept/retain\nthe client?", office],
    ]
    tbl = make_table(document, len(rows_sec1), 3, [1.2, 6.6, 9.6])
    for r, row in enumerate(rows_sec1):
        for c, val in enumerate(row):
            set_cell_text(tbl.cell(r,c), val)
    tbl.cell(0,0).merge(tbl.cell(3,0))
    tbl.cell(0,1).merge(tbl.cell(3,1))
    
    document.add_paragraph()
    add_para(document, "Section 2 – Information of the Client", size=11, bold=True)
    
    sec2_rows = [
        ["1. Individual", ""],
        ["a)\nNational Registration Identity Card\n('NRIC') or passport number", "N/A"],
        ["b) Residential and mailing address", "N/A"],
        ["2. Legal person/Legal arrangement/Club, Society and Charity", ""],
        ["a) Business registration no.", biz_reg_no],
        ["b) Address of principal place of business", address],
        ["c)\nNames of relevant persons having a\nsenior management position", senior_mgmt],
    ]
    tbl = make_table(document, len(sec2_rows), 2, [7.5, 10.0])
    for r,row in enumerate(sec2_rows):
        set_cell_text(tbl.cell(r,0), row[0])
        set_cell_text(tbl.cell(r,1), row[1])
    
    add_para(document, "Section 3 – Information of the Beneficial Owner(s)", size=11, bold=True)
    
    document.add_page_break()
    
    # ========== PAGE 4 ==========
    add_para(document, "KYC Assessment Form (continued)", size=14, bold=True)
    add_para(document, "(For Non-Gazetted Activities)", size=9)
    add_para(document, "Part A – Client Due Diligence", size=12, bold=True)
    
    bo_identity = [
        ["1. Identity of natural person(s) who ultimately has controlling ownership or power over the client:", ""],
        [NBSP, find_excel_value(kv, "Ownership Identity Summary", default="-")],
    ]
    tbl = make_table(document, len(bo_identity), 2, [7.5, 10.0])
    for r,row in enumerate(bo_identity):
        set_cell_text(tbl.cell(r,0), row[0] if r==0 else NBSP)
        set_cell_text(tbl.cell(r,1), row[1])
    
    document.add_paragraph()
    
    bo_info = [
        ["2. Information of the identified beneficial owner(s):", ""],
        ["a) Name", find_excel_value(kv, "BO Name", default="-")],
        ["b)\nNational Registration Identity Card\n('NRIC') or passport number", find_excel_value(kv, "BO NRIC", default="-")],
        ["c) Nationality", find_excel_value(kv, "BO Nationality", default="-")],
    ]
    tbl = make_table(document, len(bo_info), 2, [7.5, 10.0])
    for r,row in enumerate(bo_info):
        set_cell_text(tbl.cell(r,0), row[0])
        set_cell_text(tbl.cell(r,1), row[1])
    
    document.add_paragraph()
    add_para(document, "Section 4 - Client Risk Profiling", size=11, bold=True)
    add_para(document, "1. Based on the Artemis screening, is the client, BO and/or Director(s) classified as PEP(s) or relatives or close associates ('RCA') of PEP?", size=9.5)
    
    pep_tbl_data = [
        ["", "Yes/No", "Name of the PEP(s)"],
        ["Foreign PEP", "No", ""],
        ["Domestic PEP", pep if pep in ("Yes","No") else "No", ""],
        ["RCA of Foreign PEP", "No", ""],
        ["RCA of Domestic PEP", "No", ""],
        ["Persons who are or have been entrusted with a prominent function by an international organisation which refers to members of senior management", "No", ""],
        ["State-owned Enterprise/\nstate-invested entity", "No", ""],
    ]
    tbl = make_table(document, len(pep_tbl_data), 3, [9.5, 2.5, 6.0])
    shade_row(tbl.rows[0])
    for r,row in enumerate(pep_tbl_data):
        for c,val in enumerate(row):
            set_cell_text(tbl.cell(r,c), (val if (c!=2 or val) else ""))
    
    document.add_page_break()
    
    # ========== PAGE 5 ==========
    add_para(document, "KYC Assessment Form (continued)", size=14, bold=True)
    add_para(document, "(For Non-Gazetted Activities)", size=9)
    add_para(document, "Part A – Client Due Diligence", size=12, bold=True)
    
    sanc_tbl = [
        ["Section 4 - Client Risk Profiling", ""],
        ["2. Does the client, BO or Directors appear on the sanctions lists from www.amlcft.bnm.gov.my and www.bnm.gov.my?", "No"],
    ]
    tbl = make_table(document, len(sanc_tbl), 2,[14.5, 3.0])
    for r,row in enumerate(sanc_tbl):
        set_cell_text(tbl.cell(r,0), row[0])
        set_cell_text(tbl.cell(r,1), row[1])
    
    document.add_paragraph()
    add_para(document, "Section 4A Client Risk", size=11, bold=True)
    
    s4a = [
        ["1. What is the structure of client's business? (Complex/Simple)", find_excel_value(kv, "Business Structure", default="Simple")],
        ["2. Is the client under investigation orders and/or listed in alert list issued by authorities?", alert_list if alert_list in ("Yes","No") else "No"],
        ["3. Is the business relationship to be carried out in an unusual manner or subject to unusual requirements?", find_excel_value(kv, "Unusual requirements", default="No")],
        ["4. Is there adverse information about the client from credible sources in the public domain?", adverse if adverse in ("Yes","No") else "No"],
        ["5. Is the Client involved in higher risk industries?", find_excel_value(kv, "Higher risk industries", default="No")],
    ]
    tbl = make_table(document, len(s4a), 2, [14.5, 3.0])
    for r,row in enumerate(s4a):
        set_cell_text(tbl.cell(r,0), row[0])
        set_cell_text(tbl.cell(r,1), row[1])
    
    document.add_paragraph()
    add_para(document, "Section 4B Geographical Risk", size=11, bold=True)
    
    s4b = [
        ["1. Is the client a resident or has close connections in increased monitoring or higher risk countries?", find_excel_value(kv, "Geo risk client", default="No")],
        ["2. Is the BO, Directors, or shareholders (including intermediate & ultimate shareholder) from sanctioned or jurisdictions under increased monitoring or higher risk countries or offshore/tax havens?", find_excel_value(kv, "Geo risk shareholders", default="No")],
    ]
    tbl = make_table(document, len(s4b), 2, [14.5, 3.0])
    for r,row in enumerate(s4b):
        set_cell_text(tbl.cell(r,0), row[0])
        set_cell_text(tbl.cell(r,1), row[1])
    
    document.add_paragraph()
    add_para(document, "Section 4C Product, service, and delivery channel risks", size=11, bold=True)
    
    s4c = [
        ["1. Would the services to be provided involve providing advice or other assistance on structuring that leads to complex structure and/or structure that obscures ownership?", find_excel_value(kv, "Complex structuring", default="No")],
        ["2. Did the Firm establish non face-to-face business relationship with the client?", find_excel_value(kv, "Non face-to-face", default="No")],
    ]
    tbl = make_table(document, len(s4c), 2, [14.5, 3.0])
    for r,row in enumerate(s4c):
        set_cell_text(tbl.cell(r,0), row[0])
        set_cell_text(tbl.cell(r,1), row[1])
    
    document.add_page_break()
    
    # ========== PAGE 6 ==========
    add_para(document, "KYC Assessment Form (continued)", size=14, bold=True)
    add_para(document, "(For Non-Gazetted Activities)", size=9)
    add_para(document, "Part A – Client Due Diligence", size=12, bold=True)
    
    s5 = [
        ["1. Based on Section 4 of this form, the KYC Risk Rating by the engagement team is:", ""],
        ["KYC Risk Rating", kyc_risk],
        ["", "If your KYC Risk Rating is High or if there are PEPs/RCA of PEPs identified, please complete the Enhanced Client Due Diligence under Part B."],
    ]
    tbl = make_table(document, len(s5), 2, [12.0, 5.5])
    for r,row in enumerate(s5):
        set_cell_text(tbl.cell(r,0), row[0] if r!=2 else NBSP)
        set_cell_text(tbl.cell(r,1), row[1])
    
    document.add_page_break()
    
    # ========== PAGE 7 ==========
    add_para(document, "KYC Assessment Form (continued)", size=14, bold=True)
    add_para(document, "(For Non-Gazetted Activities)", size=9)
    add_para(document, "Part B – Enhanced Due Diligence", size=12, bold=True)
    add_para(document, "Please complete this Part B for clients with High-Risk Ratings or if there are PEPs/RCA of PEPs identified.", size=9.5)
    
    partb = [
        ["1. What is the source of wealth or source of funds for the following individuals?\nIn the case of PEPs, both sources must be obtained.", ""],
        ["PEP(s)", find_excel_value(kv, "Part B PEP SOF/SOW", default="N/A")],
    ]
    tbl = make_table(document, len(partb), 2, [12.0, 5.5])
    for r,row in enumerate(partb):
        set_cell_text(tbl.cell(r,0), row[0])
        set_cell_text(tbl.cell(r,1), row[1])
    
    # Save to bytes
    doc_bytes = io.BytesIO()
    document.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

# ========== STREAMLIT UI ==========
def main():
    # Page config
    st.set_page_config(
        page_title="KYC Document Processor",
        page_icon="📄",
        layout="wide"
    )
    
    # Custom CSS
    st.markdown("""
        <style>
        .main {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        }
        .stApp {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        }
        .upload-box {
            background: white;
            padding: 2rem;
            border-radius: 15px;
            box-shadow: 0 10px 30px rgba(0,0,0,0.2);
            margin: 1rem 0;
        }
        h1 {
            color: white;
            text-align: center;
            font-weight: 700;
            margin-bottom: 2rem;
        }
        .success-box {
            background: #10b981;
            color: white;
            padding: 1rem;
            border-radius: 10px;
            margin: 1rem 0;
        }
        .stProgress > div > div > div > div {
            background: linear-gradient(to right, #667eea, #764ba2);
        }
        </style>
    """, unsafe_allow_html=True)
    
    # Title
    st.markdown("<h1>📄 KYC Document Processor</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: white; font-size: 1.2rem; margin-bottom: 2rem;'>Transform your KYC documents into structured data - Complete 7-page Word document generation</p>", unsafe_allow_html=True)
    
    # Main container
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        with st.container():
            st.markdown("<div class='upload-box'>", unsafe_allow_html=True)
            
            st.markdown("### 📤 Upload Documents")
            
            # File uploaders
            pdf_original = st.file_uploader(
                "Upload Original PDF",
                type=['pdf'],
                help="Upload the original Tomei PDF document",
                key="pdf_original"
            )
            
            pdf_artemis = st.file_uploader(
                "Upload Artemis PDF",
                type=['pdf'],
                help="Upload the Artemis screening PDF document",
                key="pdf_artemis"
            )
            
            excel_file = st.file_uploader(
                "Upload Preliminary Excel",
                type=['xlsx', 'xls'],
                help="Upload the preliminary KYC form Excel file",
                key="excel_file"
            )
            
            st.markdown("</div>", unsafe_allow_html=True)
            
            # Process button
            if pdf_original and pdf_artemis and excel_file:
                if st.button("🚀 Process Documents", use_container_width=True, type="primary"):
                    with st.spinner("Processing documents..."):
                        try:
                            # Progress bar
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            
                            # Step 1: Convert PDFs
                            status_text.text("📄 Step 1/5: Converting Original PDF to Markdown...")
                            progress_bar.progress(10)
                            
                            md1, name1 = convert_pdf_to_markdown(pdf_original.read(), pdf_original.name)
                            progress_bar.progress(20)
                            
                            status_text.text("📄 Step 2/5: Converting Artemis PDF to Markdown...")
                            md2, name2 = convert_pdf_to_markdown(pdf_artemis.read(), pdf_artemis.name)
                            progress_bar.progress(35)
                            
                            # Step 2: Extract data
                            status_text.text("🔍 Step 3/5: Extracting Customer Profile data...")
                            customer_data = extract_customer_profile(md1)
                            progress_bar.progress(55)
                            
                            status_text.text("🔍 Step 4/5: Extracting Artemis Profile data...")
                            artemis_data = extract_artemis_profile(md2)
                            progress_bar.progress(70)
                            
                            # Step 3: Generate Word doc
                            status_text.text("📝 Step 5/5: Generating Complete 7-page Word Document...")
                            doc_bytes = generate_word_document(customer_data, artemis_data, excel_file.read())
                            progress_bar.progress(100)
                            
                            status_text.empty()
                            progress_bar.empty()
                            
                            # Success message
                            st.success("✅ Processing completed successfully! Complete 7-page KYC document generated.")
                            
                            # Download buttons
                            st.markdown("### 📥 Download Results")
                            
                            col_a, col_b, col_c = st.columns(3)
                            
                            with col_a:
                                st.download_button(
                                    label="📋 Customer JSON",
                                    data=json.dumps(customer_data, indent=2, ensure_ascii=False),
                                    file_name=f"{name1}_customer_profile.json",
                                    mime="application/json",
                                    use_container_width=True
                                )
                            
                            with col_b:
                                st.download_button(
                                    label="📋 Artemis JSON",
                                    data=json.dumps(artemis_data, indent=2, ensure_ascii=False),
                                    file_name=f"{name2}_artemis_profile.json",
                                    mime="application/json",
                                    use_container_width=True
                                )
                            
                            with col_c:
                                st.download_button(
                                    label="📝 Word Document (7 Pages)",
                                    data=doc_bytes.getvalue(),
                                    file_name=f"KYC_Form_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True
                                )
                            
                            # Document contents info
                            st.info("""
                            📄 **Generated Document Contents:**
                            - Page 1: KYC Assessment Form Summary
                            - Page 2: Client Due Diligence Approval Chain
                            - Page 3: Engagement & Client Information
                            - Page 4: Beneficial Owner Information & Risk Profiling
                            - Page 5: Client Risk Assessment (4A, 4B, 4C)
                            - Page 6: KYC Risk Rating Summary
                            - Page 7: Enhanced Due Diligence (Part B)
                            """)
                            
                            # Show preview
                            with st.expander("👁️ Preview Extracted Data"):
                                tab1, tab2 = st.tabs(["Customer Profile", "Artemis Profile"])
                                
                                with tab1:
                                    st.json(customer_data)
                                
                                with tab2:
                                    st.json(artemis_data)
                        
                        except Exception as e:
                            st.error(f"❌ Error processing documents: {str(e)}")
                            st.exception(e)
            else:
                st.info("👆 Please upload all three files to begin processing")
    
    # Footer
    st.markdown("<hr style='border: 1px solid rgba(255,255,255,0.2); margin: 3rem 0;'>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: white; opacity: 0.8;'>KYC Document Processor © 2025 | Complete 7-Page Document Generation</p>", unsafe_allow_html=True)

if __name__ == "__main__":
    main()
